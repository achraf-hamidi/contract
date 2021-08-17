import docx
import xlrd
import datetime


# print(doc.paragraphs[3].text)
# doc.paragraphs[3].text = 'new text'
# print(doc.paragraphs[3].text)
# doc.save('cd1.docx')
# i = 0 test 
# for x in doc.paragraphs[35].runs: achraf 

#     print(i)
#     print(x.text)
#     i = i+1






################################

pat = 'chauffeur.xlsx'
templ = 'CDD temp.docx'
doc = docx.Document(templ)

# To open Workbook
wb = xlrd.open_workbook(pat)
sheet = wb.sheet_by_index(0)

nbrows = sheet.nrows-3
nbcols = sheet.ncols-1

# for y in range(2):
#     yy = y+1
#     empl.append(sheet.cell_value(row, yy))
#     print(empl)
#     print(sheet.cell_value(row, yy))


def addword(inde):

    mdeltt = datetime.timedelta(days=90)
    deltt = datetime.timedelta(days=90)
    doc.paragraphs[14].runs[3].text = str(sheet.cell_value(inde, 1))
    doc.paragraphs[15].runs[2].text = str(sheet.cell_value(inde, 9))
    doc.paragraphs[16].runs[2].text = str(sheet.cell_value(inde, 6))
    doc.paragraphs[17].runs[3].text = str(sheet.cell_value(inde, 10))
    doc.paragraphs[28].runs[2].text = str(sheet.cell_value(inde, 1))
    doc.paragraphs[31].runs[1].text = str(sheet.cell_value(inde, 1))
    doc.paragraphs[35].runs[2].text = str(sheet.cell_value(inde, 3))
    doc.paragraphs[71].runs[1].text = str(sheet.cell_value(inde, 1))
    doc.paragraphs[35].runs[11].text = str(deltt)
    doc.paragraphs[70].runs[8].text = str(sheet.cell_value(inde, 1))

    doc.save('./doc/'+str(sheet.cell_value(inde, 1))+' cdd.docx')


for x in range(nbrows):
    row = 3+x
    addword(row)
